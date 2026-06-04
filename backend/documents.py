"""
MediTrace — Document Upload + Processing Router
Covers: multi-format upload (1 mark) + embedding + vector index (1 mark)
"""
import io
import csv
import json
import uuid
import time
import logging
from datetime import datetime
from typing import Optional

import fitz                         # PyMuPDF — PDF
from docx import Document as DocxDocument   # python-docx — DOCX
from fastapi import (
    APIRouter, Depends, HTTPException, UploadFile, File,
    BackgroundTasks, Form
)
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete

from auth import get_current_user, require_role
from database import get_db
from models import User, Document, DocumentChunk, Embedding
from embeddings import embed_chunks, cosine_similarity_np

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/documents", tags=["Documents"])

ALLOWED_TYPES = {"pdf", "docx", "txt", "csv"}
MAX_FILE_SIZE = 10 * 1024 * 1024   # 10 MB
CHUNK_SIZE = 750
CHUNK_OVERLAP = 100


# ─────────────────────────────────────────────
#  TEXT EXTRACTION
# ─────────────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF, preserving page boundaries."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        if text.strip():
            pages.append(f"[Page {page_num}]\n{text}")
    return "\n\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_from_csv(file_bytes: bytes) -> str:
    """Convert CSV rows to readable key: value pairs."""
    content = file_bytes.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(content))
    rows = []
    for row in reader:
        line = ", ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
        rows.append(line)
    return "\n".join(rows)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to correct extractor based on file type."""
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        return extract_text_from_docx(file_bytes)
    elif file_type == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    elif file_type == "csv":
        return extract_text_from_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def clean_text(text: str) -> str:
    """Normalize whitespace and unicode."""
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    lines = [line.strip() for line in text.splitlines()]
    # Remove excessive blank lines
    cleaned = []
    blank_count = 0
    for line in lines:
        if not line:
            blank_count += 1
            if blank_count <= 2:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)
    return "\n".join(cleaned)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[dict]:
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    chunk_index = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]
        if chunk.strip():
            chunks.append({
                "chunk_index": chunk_index,
                "chunk_text": chunk,
                "chunk_size": len(chunk),
                "start_char": start,
                "end_char": end,
            })
            chunk_index += 1
        start += chunk_size - overlap
    return chunks


# ─────────────────────────────────────────────
#  BACKGROUND PROCESSING PIPELINE
# ─────────────────────────────────────────────
async def process_document(doc_id: str, file_bytes: bytes, file_type: str):
    """
    Background task: chunk → embed → store
    Steps 4–6 of the upload pipeline
    """
    from database import AsyncSessionLocal
    async with AsyncSessionLocal() as db:
        try:
            # Fetch document
            result = await db.execute(select(Document).where(Document.id == doc_id))
            doc = result.scalar_one_or_none()
            if not doc:
                return

            doc.status = "processing"
            await db.commit()

            # Extract + clean text
            raw_text = extract_text(file_bytes, file_type)
            clean = clean_text(raw_text)
            doc.raw_text = clean

            # Chunk
            chunks_data = chunk_text(clean)
            doc.chunk_count = len(chunks_data)

            # Persist chunks
            chunk_objects = []
            for c in chunks_data:
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=doc_id,
                    **c,
                )
                db.add(chunk)
                chunk_objects.append(chunk)

            await db.flush()  # get chunk IDs

            # Embed all chunks
            texts = [c["chunk_text"] for c in chunks_data]
            vectors = embed_chunks(texts)   # returns list of 384-dim float lists

            for chunk_obj, vector in zip(chunk_objects, vectors):
                emb = Embedding(
                    id=str(uuid.uuid4()),
                    chunk_id=chunk_obj.id,
                    document_id=doc_id,
                    user_id=doc.user_id,
                    vector_json=json.dumps(vector),
                )
                db.add(emb)

            doc.status = "ready"
            await db.commit()
            logger.info(f"Document {doc_id} processed: {len(chunks_data)} chunks embedded.")

        except Exception as e:
            logger.error(f"Error processing document {doc_id}: {e}")
            async with AsyncSessionLocal() as err_db:
                result = await err_db.execute(select(Document).where(Document.id == doc_id))
                doc = result.scalar_one_or_none()
                if doc:
                    doc.status = "error"
                    await err_db.commit()


# ─────────────────────────────────────────────
#  UPLOAD ENDPOINT
# ─────────────────────────────────────────────
@router.post("/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    patient_id: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload and process a document (PDF, DOCX, TXT, CSV)."""
    # Validate file type
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=400,
            detail={
                "error": "File type not supported. Please upload PDF, DOCX, TXT, or CSV.",
                "code": "INVALID_FILE_TYPE",
            },
        )

    # Read and validate size
    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail={
                "error": "File exceeds 10MB limit.",
                "code": "FILE_TOO_LARGE",
            },
        )

    # Doctors can upload for a patient
    effective_patient_id = None
    if current_user.role == "doctor" and patient_id:
        effective_patient_id = patient_id
    elif current_user.role == "patient":
        effective_patient_id = current_user.id

    doc = Document(
        id=str(uuid.uuid4()),
        user_id=current_user.id,
        patient_id=effective_patient_id,
        filename=file.filename,
        file_type=ext,
        file_size=len(file_bytes),
        status="pending",
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    # Process in background
    background_tasks.add_task(process_document, doc.id, file_bytes, ext)

    return {
        "document_id": doc.id,
        "filename": file.filename,
        "status": "pending",
        "message": "Document received. Processing in background.",
    }


# ─────────────────────────────────────────────
#  LIST DOCUMENTS
# ─────────────────────────────────────────────
@router.get("/")
async def list_documents(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List documents visible to the current user (row-level security)."""
    if current_user.role == "admin":
        result = await db.execute(select(Document))
    elif current_user.role == "doctor":
        result = await db.execute(
            select(Document).where(Document.user_id == current_user.id)
        )
    else:
        # patient sees only their own docs
        result = await db.execute(
            select(Document).where(Document.patient_id == current_user.id)
        )

    docs = result.scalars().all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "file_type": d.file_type,
            "file_size": d.file_size,
            "upload_date": d.upload_date,
            "chunk_count": d.chunk_count,
            "status": d.status,
            "patient_id": d.patient_id,
        }
        for d in docs
    ]


# ─────────────────────────────────────────────
#  GET / PREVIEW / DELETE
# ─────────────────────────────────────────────
@router.get("/{doc_id}")
async def get_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = await _fetch_document_for_user(doc_id, current_user, db)
    return {
        "id": doc.id,
        "filename": doc.filename,
        "file_type": doc.file_type,
        "file_size": doc.file_size,
        "raw_text": doc.raw_text,
        "structured_json": doc.structured_json,
        "ai_summary": doc.ai_summary,
        "upload_date": doc.upload_date,
        "chunk_count": doc.chunk_count,
        "status": doc.status,
    }


@router.get("/{doc_id}/preview")
async def preview_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = await _fetch_document_for_user(doc_id, current_user, db)
    return {"preview_text": (doc.raw_text or "")[:5000]}


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = await _fetch_document_for_user(doc_id, current_user, db)
    await db.delete(doc)
    await db.commit()
    return {"message": "Document deleted successfully"}


# ─────────────────────────────────────────────
#  HELPER
# ─────────────────────────────────────────────
async def _fetch_document_for_user(doc_id: str, user: User, db: AsyncSession) -> Document:
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if user.role == "admin":
        return doc
    if user.role == "doctor" and doc.user_id == user.id:
        return doc
    if user.role == "patient" and doc.patient_id == user.id:
        return doc

    raise HTTPException(status_code=403, detail="Access denied to this document")
